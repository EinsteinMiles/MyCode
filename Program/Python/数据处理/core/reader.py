"""
统一文件读取：支持 Excel/CSV/JSON/TXT/DOCX 等格式
所有读取函数返回 pandas DataFrame
"""
import os
import json
import chardet
import pandas as pd

from config import DEFAULT_ENCODING, FALLBACK_ENCODINGS, logger
from utils.helpers import estimate_chunksize


def detect_encoding(path: str, sample_bytes: int = 10000) -> str:
    """自动检测文件编码"""
    try:
        with open(path, 'rb') as f:
            raw = f.read(sample_bytes)
        result = chardet.detect(raw)
        encoding = result.get('encoding', DEFAULT_ENCODING)
        if encoding and encoding.lower() in ('gb2312', 'gbk', 'gb18030'):
            return 'gbk'
        return encoding or DEFAULT_ENCODING
    except Exception:
        return DEFAULT_ENCODING


def detect_format(path: str) -> str:
    """识别文件格式"""
    ext = os.path.splitext(path)[1].lower()
    return ext


def read_file(path: str, **kwargs) -> pd.DataFrame:
    """统一读取单个文件，自动识别格式和编码

    Args:
        path: 文件路径
        **kwargs: 传递给底层读取函数的额外参数
            - sheet_name: Excel 工作表名 (默认第一个)
            - header: 表头行号，None=无表头，0=第1行是表头(默认)，
                      1=第2行是表头，以此类推
            - encoding: 编码 (默认自动检测)
            - chunksize: 分块大小

    Returns:
        pd.DataFrame
    """
    ext = detect_format(path)
    header = kwargs.pop('header', 0)
    logger.info(f"读取文件: {path} [{ext}] (header={'无' if header is None else f'第{header+1}行'})")

    # sheet_name is only for Excel files
    sheet_name = kwargs.pop('sheet_name', 0)

    if ext == '.csv':
        df = _read_csv(path, header=header, **kwargs)
    elif ext in ('.xlsx', '.xls'):
        df = _read_excel(path, header=header, sheet_name=sheet_name, **kwargs)
    elif ext == '.json':
        df = _read_json(path, **kwargs)
    elif ext == '.tsv':
        df = _read_csv(path, sep='\t', header=header, **kwargs)
    elif ext == '.txt':
        df = _read_txt(path, header=header, **kwargs)
    elif ext == '.docx':
        df = _read_docx(path, **kwargs)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

    # 如果无表头，生成友好的列名
    if header is None:
        df.columns = [f'列{i+1}' for i in range(len(df.columns))]
        logger.info(f"  无表头模式，自动生成 {len(df.columns)} 个列名: 列1 ~ 列{len(df.columns)}")

    return df


def read_multiple(paths: list[str], concat: bool = True,
                  **kwargs) -> pd.DataFrame | dict[str, pd.DataFrame]:
    """批量读取多个文件

    Args:
        paths: 文件路径列表
        concat: True 则合并为单个 DataFrame，False 返回 {文件名: DataFrame} 字典
        **kwargs: 传递给 read_file 的参数

    Returns:
        合并的 DataFrame 或 字典
    """
    dfs = {}
    for p in paths:
        try:
            df = read_file(p, **kwargs)
            name = os.path.splitext(os.path.basename(p))[0]
            # 添加来源文件列
            df['_source_file'] = os.path.basename(p)
            dfs[name] = df
            logger.info(f"  {name}: {df.shape[0]} 行 × {df.shape[1]} 列")
        except Exception as e:
            logger.error(f"读取失败 {p}: {e}")

    if concat and dfs:
        result = pd.concat(dfs.values(), ignore_index=True)
        logger.info(f"合并完成: {result.shape[0]} 行 × {result.shape[1]} 列")
        return result
    return dfs


def read_chunked(path: str, chunksize: int = None, **kwargs) -> pd.DataFrame:
    """分块读取大文件，合并返回

    Args:
        path: 文件路径
        chunksize: 每块行数
    """
    ext = detect_format(path)
    if chunksize is None:
        chunksize = estimate_chunksize(path, ext)

    logger.info(f"分块读取 [{chunksize}行/块]: {path}")

    if ext in ('.csv', '.tsv'):
        sep = '\t' if ext == '.tsv' else ','
        encoding = kwargs.pop('encoding', None) or detect_encoding(path)
        header = kwargs.pop('header', 0)
        chunks = pd.read_csv(
            path, sep=sep, encoding=encoding,
            chunksize=chunksize, header=header, **kwargs
        )
    elif ext in ('.xlsx', '.xls'):
        header = kwargs.pop('header', 0)
        chunks = pd.read_excel(path, chunksize=chunksize, header=header, **kwargs)
    else:
        return read_file(path, **kwargs)

    all_chunks = []
    for i, chunk in enumerate(chunks):
        all_chunks.append(chunk)
        if (i + 1) % 10 == 0:
            logger.info(f"  已读取 {(i+1) * chunksize} 行...")

    result = pd.concat(all_chunks, ignore_index=True)

    # 无表头时生成友好列名
    if header is None:
        result.columns = [f'列{i+1}' for i in range(len(result.columns))]

    logger.info(f"分块读取完成: {result.shape[0]} 行")
    return result


# --- 内部读取函数 ---

def _read_csv(path: str, **kwargs) -> pd.DataFrame:
    encoding = kwargs.pop('encoding', None) or detect_encoding(path)
    try:
        return pd.read_csv(path, encoding=encoding, **kwargs)
    except UnicodeDecodeError:
        for enc in FALLBACK_ENCODINGS:
            try:
                return pd.read_csv(path, encoding=enc, **kwargs)
            except UnicodeDecodeError:
                continue
        raise


def _read_excel(path: str, **kwargs) -> pd.DataFrame:
    # openpyxl 处理 xlsx, xlrd 处理 xls
    sheet_name = kwargs.pop('sheet_name', 0)
    engine = kwargs.pop('engine', None)
    if engine is None:
        engine = 'openpyxl' if path.endswith('.xlsx') else 'xlrd'
    return pd.read_excel(path, sheet_name=sheet_name, engine=engine, **kwargs)


def _read_json(path: str, **kwargs) -> pd.DataFrame:
    encoding = kwargs.pop('encoding', 'utf-8')
    with open(path, 'r', encoding=encoding) as f:
        data = json.load(f)
    orient = kwargs.pop('orient', None)
    if orient:
        return pd.read_json(path, orient=orient, encoding=encoding, **kwargs)
    # 自动处理常见 JSON 结构
    if isinstance(data, list):
        if len(data) == 0:
            return pd.DataFrame()
        if isinstance(data[0], dict):
            return pd.DataFrame(data)
        return pd.DataFrame({'value': data})
    if isinstance(data, dict):
        # 尝试找到数据数组
        for key in ('data', 'records', 'rows', 'items', 'results'):
            if key in data and isinstance(data[key], list):
                return pd.DataFrame(data[key])
        # 扁平化嵌套字典
        return pd.DataFrame([data])
    return pd.DataFrame()


def _read_txt(path: str, **kwargs) -> pd.DataFrame:
    """读取文本文件，按行读取"""
    encoding = kwargs.pop('encoding', None) or detect_encoding(path)
    sep = kwargs.pop('sep', None)
    header = kwargs.pop('header', 0)

    with open(path, 'r', encoding=encoding) as f:
        lines = f.readlines()

    # 去除空行
    lines = [l.strip() for l in lines if l.strip()]

    if not lines:
        return pd.DataFrame()

    if sep:
        # 有分隔符：解析为列
        rows = [l.split(sep) for l in lines]
    elif '\t' in lines[0] and len(lines[0].split('\t')) > 1:
        rows = [l.split('\t') for l in lines]
    else:
        # 纯文本，每行一条
        return pd.DataFrame({'line': lines})

    # 处理表头
    if header is None:
        return pd.DataFrame(rows)
    elif header > 0:
        # 跳过 header 之前的行作为表头之上的无效行
        skip_rows = rows[:header]
        data_rows = rows[header:]
        if data_rows:
            return pd.DataFrame(data_rows[1:], columns=data_rows[0])
        return pd.DataFrame()
    else:
        # header=0: 第一行是表头
        if len(rows) > 1:
            return pd.DataFrame(rows[1:], columns=rows[0])
        return pd.DataFrame(rows)


def _read_docx(path: str, **kwargs) -> pd.DataFrame:
    """读取 Word 文档中的文本和表格"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    doc = Document(path)

    # 提取段落文本
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 提取表格
    tables_data = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(cells)
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0])
            df['_table_index'] = i
            tables_data.append(df)

    if tables_data:
        result = pd.concat(tables_data, ignore_index=True)
        logger.info(f"从 DOCX 提取 {len(tables_data)} 个表格, {result.shape[0]} 行")
        return result
    elif paragraphs:
        return pd.DataFrame({'content': paragraphs})
    else:
        logger.warning(f"DOCX 文件中未找到可提取的内容: {path}")
        return pd.DataFrame()
